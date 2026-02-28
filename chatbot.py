import streamlit as st
import os
import pandas as pd
import plotly.express as px
from datetime import datetime
from mistralai import Mistral
from docx import Document
from pypdf import PdfReader
from io import BytesIO

# --- 1. CONFIGURATION & STYLE PRESTIGE ---
st.set_page_config(page_title="Lex Nexus | Cockpit Juridique", page_icon="⚖️", layout="wide")

st.markdown(r"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700&family=Inter:wght@300;400;600&display=swap');
    
    .stApp {
        background: linear-gradient(rgba(0, 0, 0, 0.9), rgba(0, 0, 0, 0.9)), 
                    url('https://images.unsplash.com/photo-1505664194779-8beaceb93744?q=80&w=2000');
        background-size: cover;
        background-attachment: fixed;
        color: #E0E0E0;
    }
    
    .main-header { font-family: 'Playfair Display', serif; color: #D4AF37; text-align: center; font-size: 3.5rem; margin-top: 10px; }
    .live-status { text-align: center; color: #00FF00; font-size: 0.7rem; letter-spacing: 3px; text-transform: uppercase; margin-bottom: 30px; animation: blink 2s infinite; }
    @keyframes blink { 0% { opacity: 1; } 50% { opacity: 0.4; } 100% { opacity: 1; } }

    .glass-card {
        background: rgba(255, 255, 255, 0.03);
        border: 1px solid rgba(212, 175, 55, 0.3);
        padding: 25px;
        border-radius: 15px;
        backdrop-filter: blur(10px);
        margin-bottom: 20px;
    }

    [data-testid="stFileUploadDropzone"] {
        border: 2px dashed #D4AF37 !important;
        background: rgba(212, 175, 55, 0.05) !important;
    }

    section[data-testid="stSidebar"] { background-color: rgba(7, 8, 12, 0.98) !important; border-right: 1px solid #D4AF37; }
</style>
""", unsafe_allow_html=True)

# --- 2. INITIALISATION DES SCORES & HISTORIQUE ---
if "legal_scores" not in st.session_state:
    st.session_state.legal_scores = {
        "Conformité": 74, "Risque": 45, "PI": 82, "Social": 61, "Fiscalité": 78
    }

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# --- 3. FONCTIONS TECHNIQUES ---

def read_file_content(file):
    name = file.name.lower()
    try:
        if name.endswith(".pdf"):
            return "\n".join([p.extract_text() for p in PdfReader(file).pages if p.extract_text()])
        elif name.endswith(".docx"):
            doc = Document(file)
            return "\n".join([p.text for p in doc.paragraphs])
        elif name.endswith(".txt"):
            return file.read().decode()
    except:
        return f"Erreur de lecture sur {file.name}"
    return ""

def generate_docx(content):
    doc = Document()
    doc.add_heading('LEX NEXUS - VERSION SÉCURISÉE 2026', 0)
    doc.add_paragraph(f"Date de génération : {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("-" * 30)
    doc.add_paragraph(content)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def plot_risk_radar():
    df = pd.DataFrame({
        "Critère": list(st.session_state.legal_scores.keys()),
        "Score": list(st.session_state.legal_scores.values())
    })
    fig = px.line_polar(df, r='Score', theta='Critère', line_close=True, color_discrete_sequence=['#D4AF37'])
    fig.update_polars(radialaxis=dict(visible=True, range=[0, 100], gridcolor="rgba(255,255,255,0.1)"), bgcolor="rgba(0,0,0,0)")
    fig.update_layout(paper_bgcolor="rgba(0,0,0,0)", font_color="white", margin=dict(l=40, r=40, t=20, b=20))
    return fig

# --- 4. NAVIGATION & SIDEBAR ---
client = Mistral(api_key=st.secrets["MISTRAL_API_KEY"])

with st.sidebar:
    st.markdown("<h1 style='color:#D4AF37; text-align:center;'>LEX NEXUS</h1>", unsafe_allow_html=True)
    menu = st.radio("NAVIGATION", ["🏛️ Dashboard", "🔬 Audit & Rédaction"])
    st.write("---")
    st.write(f"📅 **{datetime.now().strftime('%d/%m/%Y')}**")
    if st.button("✨ RÉINITIALISER TOUT"):
        st.session_state.clear()
        st.rerun()

# --- 5. PAGE : DASHBOARD ---
if menu == "🏛️ Dashboard":
    st.markdown('<p class="main-header">Cockpit Lex Nexus</p>', unsafe_allow_html=True)
    st.markdown('<p class="live-status">● SYSTÈME DE VEILLE CONNECTÉ — TEMPS RÉEL 2026</p>', unsafe_allow_html=True)
    
    k1, k2, k3 = st.columns(3)
    scores = st.session_state.legal_scores
    k1.metric("Santé Juridique", f"{sum(scores.values())//len(scores)}%", "+2%")
    k2.metric("Risque Détecté", f"{100 - scores.get('Risque', 50)}%", "-4%")
    k3.metric("Dossiers Audités", len(st.session_state.chat_history)//2, "Live")

    st.write("---")
    col_radar, col_news = st.columns([1.3, 1])
    
    with col_radar:
        st.markdown('<div class="glass-card"><h4>⚖️ Indice de Santé Juridique</h4></div>', unsafe_allow_html=True)
        st.plotly_chart(plot_risk_radar(), use_container_width=True)
    
    with col_news:
        st.markdown('<div class="glass-card"><h4>📢 Flux d\'actualités 2026</h4></div>', unsafe_allow_html=True)
        st.success("**Droit des Sociétés** : Réforme des statuts simplifiés.")
        st.info("**Jurisprudence** : Nouvel arrêt sur le secret des affaires.")
        st.warning("**Alerte** : Mise à jour obligatoire des CGV avant le 01/03.")

# --- 6. PAGE : AUDIT & RÉDACTION ---
elif menu == "🔬 Audit & Rédaction":
    st.markdown("<h2 style='text-align:center; color:#D4AF37;'>Expertise & Réécriture Automatique</h2>", unsafe_allow_html=True)
    
    files = st.file_uploader("📂 Glissez-déposez vos fichiers ici (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
    
    # --- CORRECTION DE L'ERREUR ICI ---
    # On utilise enumerate pour avoir un index 'i' unique pour chaque message
    for i, msg in enumerate(st.session_state.chat_history):
        with st.chat_message(msg["role"], avatar="⚖️" if msg["role"]=="assistant" else "👤"):
            st.markdown(msg["content"])
            if msg["role"] == "assistant":
                # On ajoute key=f"dl_{i}" pour que chaque bouton soit unique
                st.download_button(
                    label="📥 Télécharger l'acte SÉCURISÉ (.docx)",
                    data=generate_docx(msg["content"]),
                    file_name=f"LexNexus_Securise_{datetime.now().strftime('%H%M')}_{i}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{i}"
                )

    if prompt := st.chat_input("Analysez ce contrat ou demandez une réécriture..."):
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        with st.chat_message("user"): st.markdown(prompt)

        with st.chat_message("assistant", avatar="⚖️"):
            context = ""
            if files:
                for f in files: context += f"\n--- {f.name} ---\n{read_file_content(f)}\n"
            
            sys_prompt = (
                f"Tu es Lex Nexus, une IA juridique d'élite. Nous sommes le {datetime.now().strftime('%d/%m/%Y')}. "
                "Ta mission : Analyser les risques juridiques et REECRIRE systématiquement les clauses "
                "problématiques pour protéger le client tout en restant conforme au droit de 2026. "
                "Structure tes réponses avec des titres clairs et une synthèse finale."
            )
            
            stream = client.chat.stream(model="pixtral-12b-2409", messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": f"DOCUMENTS À ANALYSER:\n{context[:8000]}\n\nINSTRUCTION CLIENT: {prompt}"}
            ])
            
            full_res = ""
            placeholder = st.empty()
            for chunk in stream:
                content = chunk.data.choices[0].delta.content
                if content:
                    full_res += content
                    placeholder.markdown(full_res + "▌")
            placeholder.markdown(full_res)
            st.session_state.chat_history.append({"role": "assistant", "content": full_res})
            
            for k in st.session_state.legal_scores:
                st.session_state.legal_scores[k] = min(100, st.session_state.legal_scores[k] + 4)
            
            st.rerun()
